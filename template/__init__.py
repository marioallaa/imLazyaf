from docx import Document
from docx.enum import text
from docx.shared import Inches, Pt

def firstDocxPage(subject="", leason="", name="", teacher=""):
    d = Document()
    d.add_picture('HTFLogo.png', Inches(1.5))
    d.add_heading('Republika e Shqiperise', level=2)
    d.add_heading('Harry T.Fultz Institute,', level=2)
    d.add_heading('')
    d.add_heading("Subject: "+ subject, 0)
    d.add_paragraph('')
    d.add_paragraph('')
    d.add_paragraph(leason).alignment = text.WD_TAB_ALIGNMENT.CENTER
    d.add_paragraph('')
    d.add_paragraph('')
    d.add_paragraph('')
    d.add_paragraph('')
    d.add_paragraph('')
    d.add_paragraph('')
    d.add_paragraph('')
    d.add_paragraph('')
    d.add_paragraph('')
    d.add_paragraph('')
    d.add_heading("By: {},                                                                      To: {}".format(name, teacher), level=1)
    d.add_paragraph()
    d.add_page_break()

    


def putInfo():
    pass