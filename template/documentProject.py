from docx import Document
from docx.enum import text
from docx.shared import Inches
import os


class newDocxDocument():


    def __init__(self, title, subject, clas, name, teacher):
        self.title = title
        self.name = name
        self.subject = subject
        self.teacher = teacher
        self.clas = clas


    def firstPage(self):
        self.d = Document()
        self.d.add_picture('template\\HTFLogo.png', Inches(1.5))
        self.d.add_heading('Republika e Shqiperise', level=2)
        self.d.add_heading('Harry T.Fultz Institute,', level=2)
        self.d.add_heading('')
        self.d.add_heading("Title: " + self.title, 0)
        self.d.add_paragraph('')
        self.d.add_paragraph('')
        self.d.add_paragraph(self.subject).alignment = text.WD_TAB_ALIGNMENT.CENTER
        self.d.add_paragraph("Class: {}".format(self.clas)).alignment = text.WD_TAB_ALIGNMENT.CENTER
        self.d.add_paragraph('')
        self.d.add_paragraph('')
        self.d.add_paragraph('')
        self.d.add_paragraph('')
        self.d.add_paragraph('')
        self.d.add_paragraph('')
        self.d.add_paragraph('')
        self.d.add_paragraph('')
        self.d.add_paragraph('')
        self.d.add_paragraph('')
        self.d.add_heading(
            "By: {},                                                                      To: {}".format(self.name, self.teacher),
            level=1)
        self.d.add_paragraph()
        self.d.add_page_break()


    def addInfo(self, info):
        self.d.add_paragraph(info)
        self.d.add_page_break()

    def addRef(self, uri, refs):
        self.d.add_heading('Ref:', level=1)
        self.d.add_paragraph('')
        self.d.add_heading('This Information was taken from: ', level=2)
        self.d.add_paragraph(uri)
        self.d.add_paragraph('')
        self.d.add_paragraph('')
        self.d.add_paragraph('')
        self.d.add_heading('Other References: ', level=2)
        self.d.add_paragraph(refs)





    def saveDoc(self, path=None):
        if path == None:
            if not os.path.exists("documents"):
                 os.makedirs("documents")
                 self.d.save('documents\\{0}-{1}.docx'.format(self.title, self.name))
        else:
            path = "documents/" + path
            splitedPath = path.split('\\')
            if not os.path.exists(splitedPath[len(splitedPath)-1]):
                 os.makedirs(splitedPath[len(splitedPath)-1])
                 self.d.save(splitedPath[-1] + '.docx')


    def openImg(self):
        import webbrowser
        print("""  
         
         
         *** A Web Browser searching for images with this Subject is Opening now ***
                    
                         *** Copy and Paste any Image if You like *** """)
        url = "https://www.google.com/search?q={0}&source=lnms&tbm=isch&sa=X&ved=0ahUKEwig99z_ns_eAhWB_ywKHakgAQIQ_AUIDigB&biw=1662&bih=815&dpr=1.13".format([s for s in self.title.split(' ')])
        return webbrowser.open(url)

    def close(self):
        exit()
